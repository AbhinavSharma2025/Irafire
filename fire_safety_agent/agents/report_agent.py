from __future__ import annotations

import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from pipeline.graph import FireSafetyState


OUTPUT_PATH = Path("output/fire_safety_report.docx")


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
	cell.text = ""
	paragraph = cell.paragraphs[0]
	run = paragraph.add_run(str(text))
	run.bold = bold
	paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_table_header(table, headers: list[str]) -> None:
	header_row = table.rows[0]
	for index, header in enumerate(headers):
		_set_cell_text(header_row.cells[index], header, bold=True)


def _add_bullet(document: Document, text: str) -> None:
	paragraph = document.add_paragraph(style="List Bullet")
	paragraph.add_run(str(text))


def _find_hazard_class(hydraulic_design: dict[str, Any], room_number: Any) -> str:
	for entry in hydraulic_design.get("zone_summary", []) or []:
		if str(entry.get("room_number")) == str(room_number):
			return str(entry.get("hazard_class", ""))
	return ""


def report_agent(state: FireSafetyState) -> dict[str, Any]:
	output_path = OUTPUT_PATH
	output_path.parent.mkdir(parents=True, exist_ok=True)

	doc = Document()

	# Cover page
	title = doc.add_paragraph()
	title.alignment = WD_ALIGN_PARAGRAPH.CENTER
	title_run = title.add_run("Fire Safety Design Report")
	title_run.bold = True
	title_run.font.size = Pt(24)

	subtitle = doc.add_paragraph()
	subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
	subtitle_run = subtitle.add_run("Automated Multi-Agent Analysis")
	subtitle_run.italic = True
	subtitle_run.font.size = Pt(14)

	date_paragraph = doc.add_paragraph()
	date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
	date_paragraph.add_run(f"Date: {datetime.date.today().isoformat()}")

	input_paragraph = doc.add_paragraph()
	input_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
	input_paragraph.add_run(f"Input file: {state.get('input_file_path', '')}")

	doc.add_page_break()

	occupancy = state.get("occupancy_classification", {}) or {}
	compliance_report = state.get("compliance_report", {}) or {}
	hydraulic_design = state.get("hydraulic_design", {}) or {}
	rooms = state.get("system_design", []) or state.get("code_analysis", []) or []

	# Executive summary
	doc.add_heading("Executive Summary", level=1)
	doc.add_paragraph(
		f"Occupancy Group: {occupancy.get('occupancy_group', 'Unknown')}\n"
		f"Risk Level: {occupancy.get('risk_level', 'Unknown')}\n"
		f"Required Fire Rating: {occupancy.get('required_fire_rating', 'Unknown')}\n"
		f"Required Exits: {occupancy.get('required_exits', 'Unknown')}\n"
		f"Compliance Score: {compliance_report.get('compliance_score', 0)}\n"
		f"Total Rooms Analysed: {compliance_report.get('total_rooms_checked', len(rooms))}"
	)
	doc.add_paragraph(compliance_report.get("summary", "No summary available."))

	# Room analysis
	doc.add_heading("Room Analysis", level=1)
	room_table = doc.add_table(rows=1, cols=8)
	room_table.style = "Table Grid"
	_add_table_header(
		room_table,
		[
			"Room No",
			"Room Type",
			"Zone",
			"Hazard Class",
			"Sprinkler",
			"Detector",
			"Extinguisher",
			"Confidence",
		],
	)

	for room in rooms:
		row = room_table.add_row().cells
		fire_requirements = room.get("fire_requirements", {}) or {}
		system_design = room.get("system_design", {}) or {}
		_set_cell_text(row[0], room.get("room_number", ""))
		_set_cell_text(row[1], room.get("room_type", ""))
		_set_cell_text(row[2], room.get("zone_type", ""))
		_set_cell_text(row[3], _find_hazard_class(hydraulic_design, room.get("room_number")))
		_set_cell_text(row[4], system_design.get("sprinkler_system", ""))
		_set_cell_text(row[5], system_design.get("detection_system", ""))
		_set_cell_text(row[6], system_design.get("extinguisher_type", ""))
		_set_cell_text(row[7], fire_requirements.get("classification_confidence", ""))

	# Hydraulic calculations
	doc.add_heading("Hydraulic Design", level=1)
	doc.add_heading("Building Summary", level=2)
	building_summary = hydraulic_design.get("building_summary", {}) or {}
	doc.add_paragraph(
		f"Dominant hazard class: {building_summary.get('dominant_hazard_class', 'Unknown')}\n"
		f"Design density: {hydraulic_design.get('design_area', {}).get('design_density_lpm_m2', 'Unknown')} lpm/m2\n"
		f"Coverage per head: {hydraulic_design.get('sprinkler_design', {}).get('coverage_m2_per_head', 'Unknown')} m2/head"
	)

	doc.add_heading("Water Demand", level=2)
	water_demand = hydraulic_design.get("water_demand", {}) or {}
	water_table = doc.add_table(rows=1, cols=3)
	water_table.style = "Table Grid"
	_add_table_header(water_table, ["Sprinkler Demand", "Hose Stream", "Total Demand"])
	water_row = water_table.add_row().cells
	_set_cell_text(water_row[0], water_demand.get("sprinkler_demand_lpm", ""))
	_set_cell_text(water_row[1], water_demand.get("hose_stream_lpm", ""))
	_set_cell_text(water_row[2], water_demand.get("total_demand_lpm", ""))

	doc.add_heading("Pressure", level=2)
	pressure = hydraulic_design.get("pressure", {}) or {}
	pressure_table = doc.add_table(rows=1, cols=4)
	pressure_table.style = "Table Grid"
	_add_table_header(pressure_table, ["Required", "Available", "Margin", "Pump Required"])
	pressure_row = pressure_table.add_row().cells
	_set_cell_text(pressure_row[0], pressure.get("required_pressure_bar", ""))
	_set_cell_text(pressure_row[1], pressure.get("available_municipal_pressure_bar", ""))
	_set_cell_text(pressure_row[2], pressure.get("margin_bar", ""))
	_set_cell_text(pressure_row[3], pressure.get("pump_required", ""))

	doc.add_heading("Water Storage", level=2)
	water_storage = hydraulic_design.get("water_storage", {}) or {}
	doc.add_paragraph(
		f"Underground sump: {water_storage.get('underground_sump_liters', 0)} liters\n"
		f"Terrace tank: {water_storage.get('terrace_tank_liters', 0)} liters\n"
		f"Wet riser required: {water_storage.get('wet_riser_required', False)}"
	)

	# Compliance report
	doc.add_heading("Compliance Report", level=1)
	compliance_score_paragraph = doc.add_paragraph()
	compliance_score_run = compliance_score_paragraph.add_run(f"Compliance Score: {compliance_report.get('compliance_score', 0)}")
	compliance_score_run.bold = True

	violations = compliance_report.get("violations", []) or []
	if violations:
		violations_table = doc.add_table(rows=1, cols=5)
		violations_table.style = "Table Grid"
		_add_table_header(violations_table, ["ID", "Room", "Severity", "Description", "NBC Reference"])
		for violation in violations:
			row = violations_table.add_row().cells
			_set_cell_text(row[0], violation.get("violation_id", ""))
			_set_cell_text(row[1], violation.get("room_number", ""))
			_set_cell_text(row[2], violation.get("severity", ""))
			_set_cell_text(row[3], violation.get("description", ""))
			_set_cell_text(row[4], violation.get("nbc_reference", ""))
	else:
		doc.add_paragraph("No violations found.")

	# Review notes
	doc.add_heading("Review Notes & Assumptions", level=1)
	review_notes = hydraulic_design.get("review_notes", []) or []
	if review_notes:
		for note in review_notes:
			_add_bullet(doc, note)
	else:
		_add_bullet(doc, "No review notes provided.")

	doc.save(output_path)

	return {
		"current_agent": "report_agent",
		"report_path": str(output_path),
	}
